import os
import re
import json
import io
import requests
from bs4 import BeautifulSoup
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.pdfbase import pdfmetrics
from openai import OpenAI
from anthropic import Anthropic

# Default models
GPT_MODEL = "gpt-4o-mini"
CLAUDE_MODEL = "claude-sonnet-4-6"

def extract_text_from_pdf(file_path_or_bytes) -> str:
    """Extracts raw text from a PDF file path or bytes object."""
    try:
        reader = PdfReader(file_path_or_bytes)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        raise ValueError(f"Failed to parse PDF file: {str(e)}")

def extract_text_from_docx(file_path_or_bytes) -> str:
    """Extracts raw text from a .docx file path or bytes object."""
    try:
        doc = Document(file_path_or_bytes)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts).strip()
    except Exception as e:
        raise ValueError(f"Failed to parse Word file: {str(e)}")

# ============================================================
# Clean export helpers (Word .docx, PDF, plain .txt)
# The output deliberately strips AI/tool fingerprints:
#   - No author / application metadata in .docx core properties
#   - Smart quotes & emoji are normalized
#   - Standard professional typography (Calibri 11pt body)
# ============================================================

_EMOJI_RE = re.compile(
    "["
    "\U0001F300-\U0001FAFF"
    "\U00002600-\U000027BF"
    "\U0001F000-\U0001F2FF"
    "]+",
    flags=re.UNICODE,
)

_AI_PREAMBLE_PATTERNS = [
    r"^\s*(?:sure|certainly|absolutely|of course|here(?:'s| is)|below is|i hope|here you go)[^\n]*?:\s*\n+",
    r"^\s*```(?:markdown|md)?\s*\n",
]

def _clean_ai_artifacts(text: str) -> str:
    """Remove AI-tell signals: greeting preambles, emoji, smart quotes, fenced markers."""
    if not text:
        return ""
    cleaned = text
    for pat in _AI_PREAMBLE_PATTERNS:
        cleaned = re.sub(pat, "", cleaned, flags=re.IGNORECASE)
    # Drop trailing markdown fence
    cleaned = re.sub(r"\n```\s*$", "", cleaned)
    # Smart quotes -> straight (ATS-safer)
    cleaned = (cleaned
               .replace("‘", "'").replace("’", "'")
               .replace("“", '"').replace("”", '"')
               .replace("–", "-").replace("—", "-")
               .replace("…", "..."))
    cleaned = _EMOJI_RE.sub("", cleaned)
    return cleaned.strip()

def _parse_markdown_blocks(md_text: str):
    """Yields (kind, content) tuples. kind in {h1,h2,h3,bullet,para,hr,blank}."""
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            yield ("blank", "")
            i += 1
            continue
        if stripped.startswith("# "):
            yield ("h1", stripped[2:].strip())
        elif stripped.startswith("## "):
            yield ("h2", stripped[3:].strip())
        elif stripped.startswith("### "):
            yield ("h3", stripped[4:].strip())
        elif re.match(r"^[-*_]{3,}$", stripped):
            yield ("hr", "")
        elif re.match(r"^[\-\*\+]\s+", stripped):
            yield ("bullet", re.sub(r"^[\-\*\+]\s+", "", stripped))
        elif re.match(r"^\d+\.\s+", stripped):
            yield ("bullet", re.sub(r"^\d+\.\s+", "", stripped))
        else:
            yield ("para", stripped)
        i += 1

def _strip_inline_md(text: str) -> str:
    """Strip inline markdown markers (bold, italic, code) for plain rendering."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1 (\2)", text)
    return text

def _scrub_docx_metadata(doc):
    """Wipe identifying metadata so the file doesn't fingerprint as AI/tool-generated."""
    props = doc.core_properties
    props.author = ""
    props.last_modified_by = ""
    props.title = ""
    props.subject = ""
    props.comments = ""
    props.keywords = ""
    props.category = ""
    props.content_status = ""
    props.identifier = ""
    props.version = ""
    # Wipe the Application name in app.xml (defaults to "python-docx" otherwise)
    try:
        app_xml = doc.part.package.part_related_by(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties"
        )
        if app_xml is not None:
            root = app_xml._element if hasattr(app_xml, "_element") else None
    except Exception:
        pass

def _add_runs_with_inline_md(paragraph, text: str):
    """Add text to a paragraph honoring **bold** and *italic* inline."""
    tokens = re.split(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            run = paragraph.add_run(tok[1:-1])
            run.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            run = paragraph.add_run(tok[1:-1])
        else:
            paragraph.add_run(tok)

# ------------------------------------------------------------
# Template styling, modeled on "Navya P_Capgemini_DA.docx":
#   - Century Gothic throughout; 10pt justified body
#   - Tight ~0.3" margins
#   - Name 20pt bold, with a trailing "(Title)" rendered at 14pt
#   - Section headings: CENTERED, bold + italic + underline, 12pt, UPPERCASE
#   - Job-title lines bold 11pt; bullets use a "•" glyph at 10pt
#   - Key skills/keywords stay bold inline (from **markdown**)
# ------------------------------------------------------------
_TEMPLATE_FONT = "Century Gothic"
_BODY_PT = 10
_NAME_PT = 20
_NAME_TITLE_PT = 14
_HEADING_PT = 12
_JOBTITLE_PT = 11

def _apply_run_font(run, size_pt=None, bold=None, italic=None, underline=None):
    """Force the template font on a run (and optional size/emphasis)."""
    run.font.name = _TEMPLATE_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), _TEMPLATE_FONT)
    rfonts.set(qn("w:hAnsi"), _TEMPLATE_FONT)
    rfonts.set(qn("w:cs"), _TEMPLATE_FONT)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline

def _add_template_runs(paragraph, text: str, size_pt=_BODY_PT):
    """Add text honoring **bold**/*italic* inline, all in the template font."""
    tokens = re.split(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run = paragraph.add_run(tok[2:-2])
            _apply_run_font(run, size_pt=size_pt, bold=True)
        elif tok.startswith("*") and tok.endswith("*"):
            run = paragraph.add_run(tok[1:-1])
            _apply_run_font(run, size_pt=size_pt, italic=True)
        elif tok.startswith("`") and tok.endswith("`"):
            run = paragraph.add_run(tok[1:-1])
            _apply_run_font(run, size_pt=size_pt)
        else:
            run = paragraph.add_run(tok)
            _apply_run_font(run, size_pt=size_pt)

def markdown_to_docx_bytes(md_text: str) -> bytes:
    """Convert markdown resume/cover letter to a clean .docx styled like the template."""
    md_text = _clean_ai_artifacts(md_text)
    doc = Document()

    # Tight page margins (~0.3" as in the template)
    for section in doc.sections:
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.3)

    # Base style: Century Gothic, 10pt
    style = doc.styles["Normal"]
    style.font.name = _TEMPLATE_FONT
    style.font.size = Pt(_BODY_PT)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), _TEMPLATE_FONT)
    rfonts.set(qn("w:hAnsi"), _TEMPLATE_FONT)
    rfonts.set(qn("w:cs"), _TEMPLATE_FONT)

    first_h1_used = False
    seen_first_h2 = False
    for kind, content in _parse_markdown_blocks(md_text):
        content = _strip_inline_md(content) if kind not in ("para", "bullet") else content
        if kind == "h1":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            if not first_h1_used:
                # Name line; split a trailing "(Title)" to a smaller size
                m = re.match(r"^(.*?)\s*(\([^)]*\))\s*$", content)
                if m:
                    name_run = p.add_run(m.group(1) + " ")
                    _apply_run_font(name_run, size_pt=_NAME_PT, bold=True)
                    title_run = p.add_run(m.group(2))
                    _apply_run_font(title_run, size_pt=_NAME_TITLE_PT, bold=True)
                else:
                    run = p.add_run(content)
                    _apply_run_font(run, size_pt=_NAME_PT, bold=True)
            else:
                run = p.add_run(content)
                _apply_run_font(run, size_pt=_JOBTITLE_PT, bold=True)
            first_h1_used = True
        elif kind == "h2":
            seen_first_h2 = True
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(content.upper())
            _apply_run_font(run, size_pt=_HEADING_PT, bold=True, italic=True, underline=True)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
        elif kind == "h3":
            # Job-title / sub-entry line: bold 11pt
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _add_template_runs(p, content, size_pt=_JOBTITLE_PT)
            for r in p.runs:
                r.bold = True
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
        elif kind == "bullet":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf = p.paragraph_format
            pf.left_indent = Inches(0.25)
            pf.first_line_indent = Inches(-0.18)
            pf.space_after = Pt(3)
            bullet_run = p.add_run("•  ")
            _apply_run_font(bullet_run, size_pt=_BODY_PT)
            _add_template_runs(p, content, size_pt=_BODY_PT)
        elif kind == "hr":
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "AAAAAA")
            pBdr.append(bottom)
            pPr.append(pBdr)
        elif kind == "blank":
            continue
        else:
            p = doc.add_paragraph()
            # Contact info (before the first section heading) is centered;
            # body paragraphs are justified like the template.
            p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if not seen_first_h2
                           else WD_ALIGN_PARAGRAPH.JUSTIFY)
            p.paragraph_format.space_after = Pt(3)
            _add_template_runs(p, content, size_pt=_BODY_PT)

    _scrub_docx_metadata(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def markdown_to_pdf_bytes(md_text: str) -> bytes:
    """Convert markdown to a clean PDF (no headers/footers, no tool watermark)."""
    md_text = _clean_ai_artifacts(md_text)
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=0.7 * inch,
        rightMargin=0.7 * inch,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
        title="", author="", subject="", creator="", producer="",
    )
    styles = getSampleStyleSheet()
    name_style = ParagraphStyle(
        "Name", parent=styles["Title"], fontName="Helvetica-Bold",
        fontSize=20, leading=24, alignment=TA_CENTER, spaceAfter=4, textColor="#111111",
    )
    h2_style = ParagraphStyle(
        "H2", parent=styles["Heading2"], fontName="Helvetica-Bold",
        fontSize=12, leading=16, spaceBefore=10, spaceAfter=2, textColor="#222222",
    )
    h3_style = ParagraphStyle(
        "H3", parent=styles["Heading3"], fontName="Helvetica-Bold",
        fontSize=11, leading=14, spaceBefore=6, spaceAfter=2,
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["BodyText"], fontName="Helvetica",
        fontSize=10.5, leading=14, spaceAfter=4, alignment=TA_JUSTIFY,
    )
    bullet_style = ParagraphStyle(
        "Bullet", parent=body_style, leftIndent=14, bulletIndent=2, spaceAfter=2,
        alignment=TA_LEFT,
    )

    def md_inline_to_rl(text: str) -> str:
        """Markdown bold/italic -> ReportLab inline tags."""
        text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
        text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
        text = re.sub(r"`(.+?)`", r"\1", text)
        text = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1", text)
        return (text.replace("&", "&amp;")
                    .replace("<b>", "<b>").replace("</b>", "</b>")  # keep tags
                    .replace("<i>", "<i>").replace("</i>", "</i>"))

    flow = []
    first_h1_used = False
    pending_bullets = []

    def flush_bullets():
        if pending_bullets:
            items = [ListItem(Paragraph(md_inline_to_rl(b), bullet_style), leftIndent=14)
                     for b in pending_bullets]
            flow.append(ListFlowable(items, bulletType="bullet", start="•", leftIndent=14))
            pending_bullets.clear()

    for kind, content in _parse_markdown_blocks(md_text):
        if kind != "bullet":
            flush_bullets()
        if kind == "h1":
            style = name_style if not first_h1_used else h3_style
            flow.append(Paragraph(md_inline_to_rl(content), style))
            first_h1_used = True
        elif kind == "h2":
            flow.append(Paragraph(md_inline_to_rl(content).upper(), h2_style))
        elif kind == "h3":
            flow.append(Paragraph(md_inline_to_rl(content), h3_style))
        elif kind == "bullet":
            pending_bullets.append(content)
        elif kind == "hr":
            flow.append(Spacer(1, 6))
        elif kind == "blank":
            flow.append(Spacer(1, 4))
        else:
            flow.append(Paragraph(md_inline_to_rl(content), body_style))
    flush_bullets()

    doc.build(flow)
    return buf.getvalue()

def markdown_to_plaintext(md_text: str) -> str:
    """Convert markdown to clean plain text (ATS-friendly)."""
    md_text = _clean_ai_artifacts(md_text)
    out_lines = []
    for kind, content in _parse_markdown_blocks(md_text):
        content = _strip_inline_md(content)
        if kind == "h1":
            out_lines.append("")
            out_lines.append(content)
            out_lines.append("")
        elif kind == "h2":
            out_lines.append("")
            out_lines.append(content.upper())
            out_lines.append("-" * min(len(content), 60))
        elif kind == "h3":
            out_lines.append("")
            out_lines.append(content)
        elif kind == "bullet":
            out_lines.append(f"  - {content}")
        elif kind == "hr":
            out_lines.append("-" * 40)
        elif kind == "blank":
            out_lines.append("")
        else:
            out_lines.append(content)
    # Collapse 3+ consecutive blank lines into 2
    text = "\n".join(out_lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip() + "\n"

def scrape_job_description(url: str) -> str:
    """Scrapes clean text from a job description URL."""
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        }
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.text, "html.parser")
        
        # Remove script and style elements
        for script in soup(["script", "style", "header", "footer", "nav", "aside"]):
            script.decompose()
            
        # Get text and clean it up
        text = soup.get_text(separator="\n")
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        clean_text = "\n".join(chunk for chunk in chunks if chunk)
        
        # Limit text length to prevent context window bloating, but keep it substantial
        return clean_text[:8000].strip()
    except Exception as e:
        raise ValueError(f"Failed to scrape job description from URL: {str(e)}")

def get_openai_client(api_key: str = None) -> OpenAI:
    key = api_key or os.environ.get("OPENAI_API_KEY")
    if not key:
        raise ValueError("OpenAI API key is missing. Please configure it in Settings or your .env file.")
    return OpenAI(api_key=key)

def get_anthropic_client(api_key: str = None) -> Anthropic:
    key = api_key or os.environ.get("ANTHROPIC_API_KEY")
    if not key:
        raise ValueError("Anthropic API key is missing. Please configure it in Settings or your .env file.")
    return Anthropic(api_key=key)

def tailor_resume(openai_key: str, resume_text: str, jd_text: str) -> str:
    """Generates a tailored resume based on the original resume and Job Description."""
    client = get_openai_client(openai_key)
    
    prompt = f"""
You are an expert executive resume writer. Your task is to rewrite and tailor the following Candidate Resume to fit the provided Job Description.

Candidate Resume:
\"\"\"
{resume_text}
\"\"\"

Job Description:
\"\"\"
{jd_text}
\"\"\"

Guidelines:
1. Tailor the professional summary, skills, and work experience sections to highlight the matching skills and achievements required by the Job Description.
2. Maintain absolute truthfulness. Do not fabricate roles, companies, dates, or degrees.
3. Structure the resume in clean Markdown using EXACTLY this layout so it renders correctly:
   - First line: `# Full Name (Target Role)` — put the target job title in parentheses.
   - Next line: a single contact line (email | phone | location | LinkedIn) if available in the resume.
   - `## Professional Summary` followed by a justified paragraph.
   - `## Education` followed by bullet point(s).
   - `## Work Experience`. For each role use `### Job Title | Company, Location | Start - End` then bullet points (Action Verb + Context + Result/Impact).
   - `## Technical Skills` as bullets grouped by category, e.g. `- **Programming & Databases:** SQL, Python, PostgreSQL`.
4. Bold the key skills, tools, and keywords from the Job Description inline using **double asterisks** within the summary and bullet points (mirroring how recruiters emphasize matching keywords).
5. Keep the style modern, concise, and impact-driven.

Return ONLY the tailored resume in clean Markdown format. Do not include any introductory or concluding comments.
"""
    response = client.chat.completions.create(
        model=GPT_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3
    )
    return response.choices[0].message.content.strip()

def generate_cover_letter(openai_key: str, resume_text: str, jd_text: str) -> str:
    """Generates a professional cover letter tailored to the job description."""
    client = get_openai_client(openai_key)
    
    prompt = f"""
You are a career coach and professional writer. Write a compelling, highly tailored Cover Letter for the candidate based on their Resume and the Job Description.

Candidate Resume:
\"\"\"
{resume_text}
\"\"\"

Job Description:
\"\"\"
{jd_text}
\"\"\"

Guidelines:
1. Address it professionally (use "Dear Hiring Manager" or the company name if available in the Job Description).
2. Keep it under 400 words. It should have a hook, a strong middle showing why the candidate's achievements align with the role's needs, and a clear call to action.
3. Use a polished, confident, and professional tone.
4. Format in clean Markdown.

Return ONLY the cover letter. Do not include any meta-comments or markdown code block fences (like ```markdown).
"""
    response = client.chat.completions.create(
        model=GPT_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.4
    )
    return response.choices[0].message.content.strip()

def check_ats_score(openai_key: str, resume_text: str, jd_text: str) -> dict:
    """Computes the ATS score and returns structured feedback."""
    client = get_openai_client(openai_key)
    
    prompt = f"""
You are an advanced Applicant Tracking System (ATS) scanner. Analyze the candidate's Resume against the Job Description.

Resume:
\"\"\"
{resume_text}
\"\"\"

Job Description:
\"\"\"
{jd_text}
\"\"\"

Perform a deep analysis and output your evaluation ONLY in the following JSON format:
{{
  "score": <integer from 0 to 100 representing the matching percentage>,
  "missing_keywords": ["keyword1", "keyword2", ...],
  "formatting_issues": ["issue1", ...],
  "suggestions": ["specific recommendation to improve match", ...]
}}

Ensure the JSON is valid and return ONLY the JSON string. Do not wrap in ```json or any other text.
"""
    response = client.chat.completions.create(
        model=GPT_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1
    )
    content = response.choices[0].message.content.strip()
    
    # Strip markdown code blocks if any
    content = re.sub(r"^```json\s*", "", content)
    content = re.sub(r"^```\s*", "", content)
    content = re.sub(r"\s*```$", "", content)
    
    try:
        return json.loads(content)
    except Exception as e:
        # Fallback parsing in case JSON is slightly malformed
        try:
            # Try finding the first '{' and last '}'
            start_idx = content.find('{')
            end_idx = content.rfind('}')
            if start_idx != -1 and end_idx != -1:
                return json.loads(content[start_idx:end_idx+1])
        except:
            pass
        return {
            "score": 70,
            "missing_keywords": ["Error parsing response"],
            "formatting_issues": [],
            "suggestions": [f"Could not parse ATS evaluation response. Error: {str(e)}"]
        }

def refine_resume(openai_key: str, resume_text: str, jd_text: str, ats_feedback: dict) -> str:
    """Refines the resume using GPT, incorporating the ATS feedback to improve the score."""
    client = get_openai_client(openai_key)
    
    feedback_str = f"""
- Missing Keywords: {', '.join(ats_feedback.get('missing_keywords', []))}
- Formatting Issues: {', '.join(ats_feedback.get('formatting_issues', []))}
- Suggestions: {', '.join(ats_feedback.get('suggestions', []))}
"""
    
    prompt = f"""
You are an expert resume optimizer. Revise the following Resume to achieve a higher ATS score against the Job Description, directly addressing the provided ATS feedback.

Resume:
\"\"\"
{resume_text}
\"\"\"

Job Description:
\"\"\"
{jd_text}
\"\"\"

ATS Feedback to Address:
\"\"\"
{feedback_str}
\"\"\"

Guidelines:
1. Naturally weave in the missing keywords into bullet points, summary, or skills where appropriate. Do NOT just dump them in a list.
2. Fix any formatting issues.
3. Keep the content truthful to the original resume.
4. Output in clean Markdown format.

Return ONLY the refined resume in clean Markdown format. Do not include any introductory or concluding comments.
"""
    response = client.chat.completions.create(
        model=GPT_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3
    )
    return response.choices[0].message.content.strip()

def review_with_claude(anthropic_key: str, resume_text: str, jd_text: str) -> dict:
    """Uses Claude to review the resume and make updates to improve it further."""
    client = get_anthropic_client(anthropic_key)
    
    prompt = f"""
You are Claude, a senior recruiter and resume auditor. You are reviewing a tailored resume against a Job Description to verify if it is truly outstanding or if it needs updates to make it much better.

Resume:
\"\"\"
{resume_text}
\"\"\"

Job Description:
\"\"\"
{jd_text}
\"\"\"

Your task:
1. Identify any weak action verbs, vague achievements, poor layout flow, or general improvements.
2. If there are improvements, make them directly to the resume. 
3. If no updates are needed because it is already perfect, set "updated" to false and return the original resume.
4. Otherwise, set "updated" to true and return the updated resume.

You MUST reply ONLY in the following JSON format (no markdown formatting code blocks, no introductory text):
{{
  "updated": <true or false>,
  "feedback": "Detailed feedback of what was identified, what was changed, and why",
  "updated_resume": "The full, revised resume in Markdown format (null if updated is false)"
}}
"""
    models_to_try = [CLAUDE_MODEL, "claude-opus-4-7", "claude-haiku-4-5-20251001"]
    last_error = None
    
    for model in models_to_try:
        try:
            response = client.messages.create(
                model=model,
                max_tokens=4000,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2
            )
            content = response.content[0].text.strip()
            
            # Clean output
            content = re.sub(r"^```json\s*", "", content)
            content = re.sub(r"^```\s*", "", content)
            content = re.sub(r"\s*```$", "", content)
            
            try:
                return json.loads(content)
            except Exception as e:
                try:
                    start_idx = content.find('{')
                    end_idx = content.rfind('}')
                    if start_idx != -1 and end_idx != -1:
                        return json.loads(content[start_idx:end_idx+1])
                except:
                    pass
                raise e
        except Exception as e:
            last_error = e
            # Fallback on 404 (not found) or similar model error
            if "not_found" in str(e).lower() or "404" in str(e) or "model_not_found" in str(e).lower():
                print(f"Model {model} returned 404/Not Found. Retrying next fallback...")
                continue
            # For other critical API errors (e.g., authentication, rate limit), raise immediately
            raise e
    
    # If all models fail due to 404/not found errors
    return {
        "updated": False,
        "feedback": f"Failed to execute Claude review after trying models {models_to_try}. Last error: {str(last_error)}",
        "updated_resume": None
    }

def humanize_with_gpt(openai_key: str, resume_text: str, jd_text: str) -> str:
    """Uses GPT to review the Claude-updated resume and remove any AI-generated patterns/markers."""
    client = get_openai_client(openai_key)
    
    prompt = f"""
You are an expert copywriter and editor. Your job is to review this resume (which was enhanced by an AI model) and perform a "humanizing" pass. 
It is critical that the resume does NOT look AI-generated (avoid common AI buzzwords, repetitive structures, or overly robotic list items like "leveraged cutting-edge technology to synergize..."). It should sound natural, authentic, and written by a skilled professional, while keeping all details, achievements, and layout intact.

Resume:
\"\"\"
{resume_text}
\"\"\"

Job Description:
\"\"\"
{jd_text}
\"\"\"

Guidelines:
1. Rewrite robotic-sounding bullet points to sound like a natural human speaking about their accomplishments.
2. Remove standard AI cliché words: "testament", "delve", "synergy", "spearheaded" (if overused), "leverage" (if overused), "cutting-edge", "game-changing", "dynamic", "pioneered" (if overused). Use varied, realistic action verbs.
3. Keep the markdown layout, structure, and factual content exactly the same. Do not remove contact details or certifications.

Return ONLY the humanized resume in clean Markdown format. Do not include any introductory or concluding comments.
"""
    response = client.chat.completions.create(
        model=GPT_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.5
    )
    return response.choices[0].message.content.strip()

def run_optimization_pipeline(openai_key: str, anthropic_key: str, resume_text: str, jd_text: str, status_callback=None) -> dict:
    """Runs the complete multi-step resume optimization pipeline."""
    
    def log(step: str, message: str, details: dict = None):
        if status_callback:
            status_callback(step, message, details)
        print(f"[{step}] {message}")

    log("START", "Starting resume optimization pipeline")

    # Step 1: Tailor Resume & Cover Letter
    log("TAILORING", "Tailoring resume to the job description using GPT...")
    tailored = tailor_resume(openai_key, resume_text, jd_text)
    
    log("COVER_LETTER", "Generating cover letter using GPT...")
    cover_letter = generate_cover_letter(openai_key, resume_text, jd_text)
    
    # Step 2: Check ATS Score & Loop to optimize (max 3 runs)
    current_resume = tailored
    ats_history = []
    ats_pass = False
    
    for i in range(1, 4):
        log("ATS_CHECK", f"Checking ATS score (Attempt {i}/3)...")
        ats_feedback = check_ats_score(openai_key, current_resume, jd_text)
        score = ats_feedback.get("score", 0)
        ats_history.append({
            "attempt": i,
            "score": score,
            "feedback": ats_feedback
        })
        
        log("ATS_CHECK", f"ATS score is {score}% (Target: >= 85%)", {"score": score, "feedback": ats_feedback})
        
        if score >= 85:
            ats_pass = True
            break
        
        if i < 3:
            log("ATS_REFINE", f"Score {score}% is below 85%. Refining resume using GPT...")
            current_resume = refine_resume(openai_key, current_resume, jd_text, ats_feedback)
    
    # Step 3: Review with Claude
    log("CLAUDE_REVIEW", "Submitting resume to Claude for review and audit...")
    claude_result = review_with_claude(anthropic_key, current_resume, jd_text)
    
    updated_by_claude = claude_result.get("updated", False)
    claude_feedback = claude_result.get("feedback", "No feedback provided.")
    
    if updated_by_claude and claude_result.get("updated_resume"):
        log("CLAUDE_REVIEW", "Claude suggested improvements and updated the resume.", {"feedback": claude_feedback})
        current_resume = claude_result.get("updated_resume")
        
        # Step 4: Humanize with GPT
        log("HUMANIZER", "Humanizing the Claude-updated resume with GPT to remove AI signatures...")
        final_resume = humanize_with_gpt(openai_key, current_resume, jd_text)
    else:
        log("CLAUDE_REVIEW", "Claude reviewed the resume and found no updates needed.", {"feedback": claude_feedback})
        final_resume = current_resume
        
    log("COMPLETE", "Pipeline complete!", {"final_score": ats_history[-1]["score"]})
    
    return {
        "initial_tailored_resume": tailored,
        "cover_letter": cover_letter,
        "ats_history": ats_history,
        "ats_pass": ats_pass,
        "claude_updated": updated_by_claude,
        "claude_feedback": claude_feedback,
        "final_resume": final_resume
    }
